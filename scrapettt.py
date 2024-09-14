import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import RGBColor
import re

# Base URL
base_url = "https://talktajiktoday.com/word-a-day/page/"

# List to store all posts
all_posts = []

# Regular expression to detect Cyrillic characters
cyrillic_pattern = re.compile('[\u0400-\u04FF]+')

# Loop through all pages
for page_num in range(1, 82):  # Adjust the range as needed
    # Construct the URL for the current page
    if page_num == 1:
        url = 'https://talktajiktoday.com/word-a-day/'
    else:
        url = f'https://talktajiktoday.com/word-a-day/page/{page_num}/'
    
    # Fetch the page content
    response = requests.get(url)
    
    # Check if the request was successful
    if response.status_code == 200:
        # Parse the HTML content
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Find all posts
        posts = soup.find_all('div', class_='entry-content')
        
        # Extract and store the content of each post
        for post in posts:
            # print("post", post)
            # Extract the heading
            heading_tag = post.find('h2')
            heading = heading_tag.get_text(strip=True) if heading_tag else ""
            #print("heading", heading)
            # Remove the heading from the post content
            if heading_tag:
                heading_tag.decompose()
            # Extract the rest of the content
            # If div contains "Bonus" get the content of the div
            for div_tag in post.find_all('div'):
                if "Bonus" in div_tag.get_text():
                    bonus = div_tag.get_text()
                    #print("bonus", bonus)
                    # remove that div that contains "Bonus" from the post
                    div_tag.decompose()
                    # remove newline after "Opposite -"
                    bonus = re.sub(r'Opposite -\s+', 'Opposite - ', bonus)
            # print("post", post)
            # Extract the rest of the content
            main_content = post.get_text(separator="\n", strip=True)

            all_posts.append((heading, main_content, bonus))

# sort posts by heading
all_posts.sort(key=lambda x: x[0].lower())

# Create a new Word document
doc = Document()

# Add each post to the documentS
for heading, main_content, bonus in all_posts:
    if heading:
        doc.add_heading(heading, level=2)
    if main_content:
        paragraph = doc.add_paragraph()
        sentence_pattern = re.compile(r'([^.!?]*[.!?])')
        sentences = sentence_pattern.findall(main_content)
        for sentence in sentences:
            #print("sentence", sentence)
            run = paragraph.add_run(sentence + " ")
            if cyrillic_pattern.search(sentence):
                run.bold = True
        if bonus:
            # Add the bonus content in green color to the same paragraph
            run = paragraph.add_run("\n" + bonus)
            font = run.font
            font.color.rgb = RGBColor(0, 128, 0)
    elif bonus:
        # If there is no main content, add the bonus content as a new paragraph
        run = doc.add_paragraph().add_run(bonus)
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0)


# Save the document
doc.save("scraped-ttt-posts.docx")

# print("Posts have been saved to scraped_posts.docx")